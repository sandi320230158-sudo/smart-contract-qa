"""
Smart Contract Q&A Assistant
Run with: python app.py
"""

import os
import re
import io
import time
from pathlib import Path
from typing import List, Optional

import fitz                        # PyMuPDF
import docx as python_docx         # python-docx

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_community.llms import HuggingFacePipeline
from transformers import pipeline as hf_pipeline

import gradio as gr

from config import CONFIG

# ---------------------------------------------------------------------------
# 1. Embedding model
# ---------------------------------------------------------------------------
print(f"Loading embedding model: {CONFIG['embed_model']} …")
embeddings = HuggingFaceEmbeddings(
    model_name=CONFIG["embed_model"],
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True},
)
print("Embeddings ready.")

# ---------------------------------------------------------------------------
# 2. LLM
# ---------------------------------------------------------------------------
print(f"Loading LLM: {CONFIG['llm_model']} …")
_hf_pipe = hf_pipeline(
    "text2text-generation",
    model=CONFIG["llm_model"],
    max_new_tokens=CONFIG["max_new_tokens"],
    do_sample=False,
    truncation=True,
)
llm = HuggingFacePipeline(pipeline=_hf_pipe)
print("LLM ready.")

# ---------------------------------------------------------------------------
# 3. Text splitter
# ---------------------------------------------------------------------------
splitter = RecursiveCharacterTextSplitter(
    chunk_size=CONFIG["chunk_size"],
    chunk_overlap=CONFIG["chunk_overlap"],
    separators=["\n\n", "\n", ". ", " ", ""],
)

# ---------------------------------------------------------------------------
# 4. Document ingestion helpers
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes, filename: str) -> List[Document]:
    docs = []
    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text().strip()
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": filename, "page": page_num},
                ))
    print(f"  PDF: {len(docs)} pages extracted from '{filename}'")
    return docs


def extract_text_from_docx(file_bytes: bytes, filename: str) -> List[Document]:
    doc = python_docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    docs, buffer, word_count, page_num = [], [], 0, 1
    for para in paragraphs:
        buffer.append(para)
        word_count += len(para.split())
        if word_count >= 500:
            docs.append(Document(
                page_content="\n".join(buffer),
                metadata={"source": filename, "page": page_num},
            ))
            buffer, word_count, page_num = [], 0, page_num + 1
    if buffer:
        docs.append(Document(
            page_content="\n".join(buffer),
            metadata={"source": filename, "page": page_num},
        ))
    print(f"  DOCX: {len(docs)} sections extracted from '{filename}'")
    return docs


def ingest_file(file_bytes: bytes, filename: str) -> List[Document]:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes, filename)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes, filename)
    raise ValueError(f"Unsupported file type: {ext}. Upload a PDF or DOCX.")


def chunk_documents(docs: List[Document]) -> List[Document]:
    chunks = splitter.split_documents(docs)
    print(f"  Chunked {len(docs)} pages → {len(chunks)} chunks")
    return chunks

# ---------------------------------------------------------------------------
# 5. Vector store helpers
# ---------------------------------------------------------------------------

def build_vector_store(chunks: List[Document]) -> FAISS:
    return FAISS.from_documents(chunks, embeddings)


def similarity_score(query: str, vs: FAISS) -> float:
    results = vs.similarity_search_with_score(query, k=1)
    if not results:
        return 0.0
    _, l2_dist = results[0]
    return max(0.0, 1.0 - (l2_dist ** 2) / 2)

# ---------------------------------------------------------------------------
# 6. RAG chain
# ---------------------------------------------------------------------------

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


def build_qa_chain(vs: FAISS):
    retriever = vs.as_retriever(search_kwargs={"k": CONFIG["top_k"]})
    prompt = ChatPromptTemplate.from_messages([
        ("system",
         "You are a precise legal document assistant. "
         "Answer the question using ONLY the contract excerpts below. "
         "If the answer is not in the excerpts, say: "
         "'I could not find this information in the uploaded document.'\n\n"
         "Contract excerpts:\n{context}"),
        ("human", "{question}"),
    ])
    chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    return retriever, chain

# ---------------------------------------------------------------------------
# 7. Safety guard-rails
# ---------------------------------------------------------------------------
UNSAFE_PATTERNS = [
    r"ignore (previous|all) instructions?",
    r"jailbreak",
    r"forget (you are|your instructions?)",
    r"pretend (you are|to be)",
]

def is_prompt_injection(text: str) -> bool:
    return any(re.search(p, text.lower()) for p in UNSAFE_PATTERNS)

# ---------------------------------------------------------------------------
# 8. Q&A entry point
# ---------------------------------------------------------------------------

def answer_question(query, chain_tuple, vs, history):
    if is_prompt_injection(query):
        return "Guard-rail: prompt injection detected. Query blocked.", []

    score = similarity_score(query, vs)
    if score < CONFIG["relevance_threshold"]:
        return (
            f"Your question does not appear related to the uploaded document "
            f"(relevance score: {score:.2f}). Please ask about the contract.",
            [],
        )

    retriever, chain = chain_tuple
    answer = chain.invoke(query)

    source_docs = retriever.invoke(query)
    sources = []
    for doc in source_docs:
        meta = doc.metadata
        citation = f"Source: {meta.get('source', 'document')}, page {meta.get('page', '?')}"
        snippet = doc.page_content[:120].replace("\n", " ") + "..."
        sources.append(f"{citation} -- {snippet}")

    return answer, sources

# ---------------------------------------------------------------------------
# 9. Document summarisation
# ---------------------------------------------------------------------------

def summarize_document(docs: List[Document], max_chars: int = 3000) -> str:
    if not docs:
        return "No document loaded."
    partial_summaries = []
    for i, doc in enumerate(docs[:10]):
        text = doc.page_content[:800]
        prompt = f"Summarize this contract section in 2-3 sentences:\n{text}"
        out = _hf_pipe(prompt)[0]["generated_text"].strip()
        partial_summaries.append(out)
    combined = " ".join(partial_summaries)[:max_chars]
    final_prompt = (
        "You are a legal document assistant. "
        "Write a concise executive summary of this contract "
        f"based on these section summaries:\n{combined}"
    )
    return _hf_pipe(final_prompt)[0]["generated_text"].strip()

# ---------------------------------------------------------------------------
# 10. Gradio UI
# ---------------------------------------------------------------------------
vector_store: Optional[FAISS] = None
qa_chain = None
current_filename = ""
raw_docs: List[Document] = []


def process_document(file):
    global vector_store, qa_chain, current_filename, raw_docs
    if file is None:
        return "Please upload a file first."
    try:
        with open(file.name, "rb") as f:
            file_bytes = f.read()
        filename = Path(file.name).name
        raw_docs = ingest_file(file_bytes, filename)
        chunks = chunk_documents(raw_docs)
        vector_store = build_vector_store(chunks)
        qa_chain = build_qa_chain(vector_store)
        current_filename = filename
        return f"✅ '{filename}' processed — {len(raw_docs)} pages, {len(chunks)} chunks indexed."
    except Exception as e:
        return f"❌ Error: {e}"


def summarize_doc():
    if not raw_docs:
        return "No document loaded. Process a file first."
    try:
        return summarize_document(raw_docs)
    except Exception as e:
        return f"❌ Error: {e}"


def chat(query, history):
    if not query.strip():
        return history, ""
    if qa_chain is None:
        history.append({"role": "user", "content": query})
        history.append({"role": "assistant", "content": "⚠️ Upload and process a document first."})
        return history, ""
    try:
        answer, sources = answer_question(query, qa_chain, vector_store, history)
        src_text = ("\n\nSources:\n" + "\n".join(f"• {s}" for s in sources)) if sources else ""
        response = answer + src_text
    except Exception as e:
        response = f"❌ Error: {e}"
    history.append({"role": "user", "content": query})
    history.append({"role": "assistant", "content": response})
    return history, ""


with gr.Blocks(theme=gr.themes.Soft(), title="Smart Contract Q&A") as demo:
    gr.Markdown("# 📄 Smart Contract Q&A Assistant")
    gr.Markdown("Upload a PDF or DOCX contract, then ask questions about it.")

    with gr.Tabs():
        with gr.Tab("📁 Upload"):
            file_input = gr.File(label="Choose PDF or DOCX", file_types=[".pdf", ".docx"])
            process_btn = gr.Button("Process Document", variant="primary", size="lg")
            process_status = gr.Textbox(label="Status", interactive=False, lines=2)
            gr.Markdown("---")
            summarize_btn = gr.Button("Summarize Contract", variant="secondary")
            summary_out = gr.Textbox(label="Executive Summary", interactive=False, lines=8)

            process_btn.click(fn=process_document, inputs=file_input, outputs=process_status)
            summarize_btn.click(fn=summarize_doc, inputs=None, outputs=summary_out)

        with gr.Tab("💬 Chat"):
            chatbot = gr.Chatbot(height=420, label="Contract Assistant",
                                 type="messages", show_copy_button=True)
            query_box = gr.Textbox(
                placeholder="Ask a question about the contract…",
                label="Your question", lines=2,
            )
            with gr.Row():
                send_btn = gr.Button("Send", variant="primary")
                clear_btn = gr.Button("Clear Chat", variant="stop")

            send_btn.click(fn=chat, inputs=[query_box, chatbot], outputs=[chatbot, query_box])
            query_box.submit(fn=chat, inputs=[query_box, chatbot], outputs=[chatbot, query_box])
            clear_btn.click(fn=lambda: ([], ""), inputs=None, outputs=[chatbot, query_box])

if __name__ == "__main__":
    demo.launch(debug=True, share=True)
